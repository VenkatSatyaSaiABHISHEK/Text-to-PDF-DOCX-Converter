import re, os
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

def parse_text(text):
    code_blocks = re.findall(r'```(.*?)```', text, re.DOTALL)
    non_code_parts = re.split(r'```.*?```', text, flags=re.DOTALL)
    return non_code_parts, code_blocks

def style_code_block(paragraph):
    run = paragraph.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    paragraph._p.get_or_add_pPr().append(shading_elm)

def create_docx(text):
    doc = Document()
    non_code, code_blocks = parse_text(text)
    for i, part in enumerate(non_code):
        if part.strip():
            doc.add_paragraph(part.strip())
        if i < len(code_blocks):
            para = doc.add_paragraph()
            para.add_run(code_blocks[i].strip())
            style_code_block(para)
    file_path = "output.docx"
    doc.save(file_path)
    return file_path

def create_pdf(text):
    filename = "output.pdf"
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]
    code_style = ParagraphStyle(name='CodeBlock', fontName='Courier', fontSize=9, backColor=colors.whitesmoke, leftIndent=10, borderPadding=5, leading=12, alignment=TA_LEFT)
    elements = []
    non_code, code_blocks = parse_text(text)
    for i, part in enumerate(non_code):
        if part.strip():
            elements.append(Paragraph(part.strip().replace("\n", "<br/>"), normal_style))
            elements.append(Spacer(1, 12))
        if i < len(code_blocks):
            code_text = code_blocks[i].strip().replace(" ", "&nbsp;").replace("\n", "<br/>")
            elements.append(Paragraph(code_text, code_style))
            elements.append(Spacer(1, 12))
    doc.build(elements)
    return filename
